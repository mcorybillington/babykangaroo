import ntpath
import docx
import argparse
from enchant import Dict
from nltk.corpus import wordnet
from nltk.data import load


def options():
    parser = argparse.ArgumentParser(description="Sound smarter, immediately.")
    parser.add_argument('-c', '--corporate', action='store_true', help='Corporate syntax')
    parser.add_argument('docx_file', help="Document to babykangaroo-ify")
    return parser.parse_args()


def capitalize_sentences(sentences):
    sent_tokenizer = load('tokenizers/punkt/english.pickle')
    new_sentences = sent_tokenizer.tokenize(' '.join(sentences))
    new_sentences = [sent.capitalize() for sent in new_sentences]
    return ' '.join(new_sentences)


def longest(wordlist):
    return max(wordlist, key=lambda s: (len(s), s))


def babykangarooify(path, arg):
    head, tail = ntpath.split(path)
    d = Dict("en_US")
    doc = docx.Document(path)
    new_doc = docx.Document()

    for paragraph in doc.paragraphs:
        new_paragraph = []
        for word in paragraph.text.split():
            if 'joey' in word.lower():
                new_paragraph.append('baby kangaroo')
                continue
            syns = [l.name() for syn in wordnet.synsets(word) for l in syn.lemmas() if d.check(l.name())]
            if arg.corporate and syns:
                synergize_words = open('buzzwords/corporate.txt').read().splitlines()
                possible = []
                for syn_word in synergize_words:
                    synergy_syns = [l.name() for syn in wordnet.synsets(syn_word) for l in syn.lemmas() if d.check(l.name())]
                    synergy_exists = [i for i in synergy_syns if i in syns]
                    if synergy_exists:
                        possible.append(syn_word)
                new_word = max(possible, key=lambda s: (len(s), s)) if possible else word
                new_paragraph.append(new_word)
            elif syns:
                print('here')
                new_word = max(syns, key=lambda s: (len(s), s))
                new_paragraph.append(new_word)
            else:
                new_paragraph.append(word)
        new_cap_paragraph = capitalize_sentences(new_paragraph)
        new_doc.add_paragraph(new_cap_paragraph)
    new_doc.save(head + '/bk_' + tail)


def main():
    arg = options()
    babykangarooify(arg.docx_file, arg)

if __name__ == '__main__':
    main()
